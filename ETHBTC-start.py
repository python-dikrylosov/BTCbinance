import random
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import yfinance as yf
import pandas as pd
from docx.shared import Mm
# создание пустого документа
doc = Document()
# без указания аргумента `level`
real_date = str(time.strftime("%Y-%m-%d"))
# добавляется заголовок "Heading 1"
head = doc.add_heading('Робот/бинанс LTCBTC ' + real_date )
# выравнивание посередине
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

#data_LTCBTC = yf.download("LTC-BTC",start="2022-09-22",end=real_date,interval="1m")
#data_LTCBTC_filter_close = data_LTCBTC.filter(["Close"])
#datapd = pd.DataFrame(data_LTCBTC_filter_close)
#datapd.to_csv(str(time.strftime("%Y-%m-%d"))+'.csv')
api_key = ""
secret_key = ""
from binance.client import Client
client = Client(api_key, secret_key)

info = client.get_all_tickers()
symbol_ETHBTC = info[0]
symbol_ETHBTC_symbol = symbol_ETHBTC["symbol"]
symbol_ETHBTC_price = symbol_ETHBTC["price"]

symbol_LTCBTC = info[1]
symbol_LTCBTC_symbol = symbol_LTCBTC["symbol"]
symbol_LTCBTC_price = symbol_LTCBTC["price"]

symbol_BNBBTC = info[2]
symbol_BNBBTC_symbol = symbol_BNBBTC["symbol"]
symbol_BNBBTC_price = symbol_BNBBTC["price"]

symbol_NEOBTC = info[3]
symbol_NEOBTC_symbol = symbol_NEOBTC["symbol"]
symbol_NEOBTC_price = symbol_NEOBTC["price"]


data_safe_close = open('ltcbtc.csv',"a")
data_safe_close.write(str(time.strftime("%Y-%m-%d %H:%M:%S+00:00")))
data_safe_close.write(",")
data_safe_close.write(str(symbol_LTCBTC_price))
data_safe_close.write(",")
data_safe_close.write("\n")
data_safe_close.close()

data = yf.download("LTC-BTC",start="2022-09-22",end=real_date,interval="1m")
data_fc = data.filter(["Close"])
print(data_fc)
plt.plot(data_fc)
plt.savefig('data_LTCBTC_filter_close.jpg')
table = doc.add_table(rows=2, cols=2)
# данные таблицы без названий колонок

for i in range(0):
    symbol = info[i]
    symbol_symbol = symbol["symbol"]
    symbol_price = symbol["price"]
    data = [i,symbol_symbol,symbol_price]
    print(data)
    data_file_safe = open("data.csv","a")
    data_file_safe.write(str(data))
    data_file_safe.write(",")
    data_file_safe.write("\n")
    data_file_safe.close()

symbol_ETHBTC = info[0]
symbol_LTCBTC = info[1]
symbol_BNBBTC = info[2]
symbol_NEOBTC = info[3]
symbol_BTCUSDT = info[11]
symbol_ETHUSDT = info[12]
symbol_BNBUSD = info[98]
symbol_LTCUSD = info[190]
symbol_XRPUSD = info[306]
symbol_BTCRUB = info[666]
symbol_BTCRUB_symbol = symbol_BTCRUB["symbol"]
symbol_BTCRUB_price = symbol_BTCRUB["price"]
symbol_DOTUSD = info[954]
symbol_AXSUSD = info[1139]
# https://python-binance.readthedocs.io/en/latest/overview.html

BTC_acount = client.get_asset_balance(asset='BTC')
#print(BTC_acount)
balance_btc = client.get_asset_balance(asset='BTC')
# print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
balance_usd_btc_usd_free_locked_sum = float(balance_btc["free"]) + float(balance_btc["locked"])
balance_usd_btc_usd_present = balance_usd_btc_usd_free_locked_sum * float(symbol_BTCUSDT["price"])

RUB_acount = client.get_asset_balance(asset='RUB')
#print(RUB_acount)
balance_RUB = client.get_asset_balance(asset='RUB')
# print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
balance_usd_RUB_usd_free_locked_sum = float(balance_RUB["free"]) + float(balance_RUB["locked"])
balance_usd_RUB_usd_present = balance_usd_RUB_usd_free_locked_sum * float(symbol_BTCUSDT["price"])

balance_ETH = client.get_asset_balance(asset='ETH')
# print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
balance_usd_ETH_usd_free_locked_sum = float(balance_ETH["free"]) + float(balance_ETH["locked"])
balance_usd_ETH_usd_present = balance_usd_ETH_usd_free_locked_sum * float(symbol_ETHBTC["price"])

balance_LTC = client.get_asset_balance(asset='LTC')
# print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
balance_usd_LTC_usd_free_locked_sum = float(balance_LTC["free"]) + float(balance_LTC["locked"])
balance_usd_LTC_usd_present = balance_usd_ETH_usd_free_locked_sum * float(symbol_LTCBTC["price"])

items = (
                (0, symbol_ETHBTC_symbol, symbol_ETHBTC_price, balance_btc["free"], balance_ETH["free"]),
                (1, symbol_LTCBTC_symbol, symbol_LTCBTC_price, balance_btc["free"], balance_LTC["free"]),
                (2, symbol_BNBBTC_symbol, symbol_BNBBTC_price, balance_btc["free"], balance_RUB["free"]),
                (3, symbol_NEOBTC_symbol, symbol_NEOBTC_price, balance_btc["free"], balance_RUB["free"]),
                (666, symbol_BTCRUB_symbol, symbol_BTCRUB_price, balance_btc["free"], balance_RUB["free"]),
            )

#print(items)

# добавляем таблицу с одной строкой
# для заполнения названий колонок
table = doc.add_table(1, len(items[0]))
# определяем стиль таблицы
table.style = 'Light Shading Accent 1'
# Получаем строку с колонками из добавленной таблицы
head_cells = table.rows[0].cells
# добавляем названия колонок
for i, item in enumerate(['Номер', 'Название', 'Курс', "свободные BTC", "свободные"]):
    p = head_cells[i].paragraphs[0]
    # название колонки
    p.add_run(item).bold = True
    # выравниваем посередине
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# добавляем данные к существующей таблице
for row in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(row):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        # если последняя ячейка
        if i == 2:
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Arial'

dataETHBTC = yf.download("ETH-BTC", start="2014-04-01", end=real_date, interval="1d")
dataETHBTC_fc = dataETHBTC.filter(["Close"])
#print(dataETHBTC_fc)
plt.plot(dataETHBTC_fc)
plt.savefig('data_ETHBTC_filter_close.jpg')


def add_img(image_path, width=None, height=None):
    """Метод Document.add_picture()"""
    img = doc.add_paragraph().add_run().add_picture(image_path, width, height)
    return img


doc.add_heading(str(dataETHBTC_fc))
add_img('data_ETHBTC_filter_close.jpg', width=Mm(100))


def add_img(image_path, width=None, height=None):
    """Метод Document.add_picture()"""
    img = doc.add_paragraph().add_run().add_picture(image_path, width, height)
    return img


doc.add_heading(str(data_fc))
add_img('data_LTCBTC_filter_close.jpg', width=Mm(100))

dataBNBBTC = yf.download("BNB-BTC", start="2014-04-01", end=real_date, interval="1d")
dataBNBBTC_fc = dataETHBTC.filter(["Close"])
#print(dataBNBBTC_fc)
plt.plot(dataBNBBTC_fc)
plt.savefig('data_BNBBTC_filter_close.jpg')


def add_img(image_path, width=None, height=None):
    """Метод Document.add_picture()"""
    img = doc.add_paragraph().add_run().add_picture(image_path, width, height)
    return img


doc.add_heading(str(dataBNBBTC_fc))
add_img('data_BNBBTC_filter_close.jpg', width=Mm(100))

dataNEOBTC = yf.download("NEO-BTC", start="2014-04-01", end=real_date, interval="1d")
dataNEOBTC_fc = dataNEOBTC.filter(["Close"])
data_safe_NEOBTC = dataNEOBTC_fc.to_csv("NEO-BTC.csv")
#print(dataNEOBTC_fc)
plt.plot(dataNEOBTC_fc)
plt.savefig('data_NEOBTC_filter_close.jpg')


def add_img(image_path, width=None, height=None):
    """Метод Document.add_picture()"""
    img = doc.add_paragraph().add_run().add_picture(image_path, width, height)
    return img


doc.add_heading(str(dataNEOBTC_fc))
add_img('data_NEOBTC_filter_close.jpg', width=Mm(100))

"""Document.add_heading() #добавляет абзац заголовка,
Document.add_page_break() #добавляет разрыв страницы,
Document.add_paragraph() #добавляет абзац,
Document.add_picture() #добавляет изображения в отдельный абзац,
Document.add_section() #добавляет новую секцию,
Document.add_table() #добавляет новую таблицу,
Document.core_properties #основные свойства документа,
Document.inline_shapes #список объектов изображений InlineShape,
Document.paragraphs #список объектов абзацев Paragraph,
Document.save() #сохраняет этот документ,
Document.sections #список объектов раздела Section,
Document.settings #объект Settings,
Document.styles #объект Styles,
Document.tables #список объектов таблиц Table,"""

doc.save(str(time.strftime("%Y-%m-%d")) + '.docx')

"""if str(balance_RUB["free"]) >= str(110) :
    order_buy_BTCRUB = client.order_limit_buy(
        symbol='BTCRUB',
        quantity=0.0001,
        price=symbol_BTCRUB_price)
    print(order_buy_BTCRUB)
elif str(balance_RUB["free"]) <= str(110):
    if balance_btc["free"] >= str(0.00011):
        order_sell_BTCRUB = client.order_limit_sell(
            symbol='BTCRUB',
            quantity=0.0001,
            price=symbol_BTCRUB_price)
        print(order_sell_BTCRUB)"""


while True:
    time.sleep(10)
    min_lovume_rub = 150
        symbol_ETHBTC = info[0]
        symbol_LTCBTC = info[1]
        symbol_BNBBTC = info[2]
        symbol_NEOBTC = info[3]
        symbol_BTCUSDT = info[11]
        symbol_ETHUSDT = info[12]
        symbol_BNBUSD = info[98]
        symbol_LTCUSD = info[190]
        symbol_XRPUSD = info[306]
        symbol_BTCRUB = info[666]
        symbol_BTCRUB_symbol = symbol_BTCRUB["symbol"]
        symbol_BTCRUB_price = symbol_BTCRUB["price"]
        symbol_DOTUSD = info[954]
        symbol_AXSUSD = info[1139]
        # https://python-binance.readthedocs.io/en/latest/overview.html

        BTC_acount = client.get_asset_balance(asset='BTC')
        # print(BTC_acount)
        balance_btc = client.get_asset_balance(asset='BTC')
        # print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
        balance_usd_btc_usd_free_locked_sum = float(balance_btc["free"]) + float(balance_btc["locked"])
        balance_usd_btc_usd_present = balance_usd_btc_usd_free_locked_sum * float(symbol_BTCUSDT["price"])

        RUB_acount = client.get_asset_balance(asset='RUB')
        # print(RUB_acount)
        balance_RUB = client.get_asset_balance(asset='RUB')
        # print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
        balance_usd_RUB_usd_free_locked_sum = float(balance_RUB["free"]) + float(balance_RUB["locked"])
        balance_usd_RUB_usd_present = balance_usd_RUB_usd_free_locked_sum * float(symbol_BTCUSDT["price"])

        balance_ETH = client.get_asset_balance(asset='ETH')
        # print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
        balance_usd_ETH_usd_free_locked_sum = float(balance_ETH["free"]) + float(balance_ETH["locked"])
        balance_usd_ETH_usd_present = balance_usd_ETH_usd_free_locked_sum * float(symbol_ETHBTC["price"])

        balance_LTC = client.get_asset_balance(asset='LTC')
        # print(balance_btc["asset"],balance_btc["free"],balance_btc["locked"])
        balance_usd_LTC_usd_free_locked_sum = float(balance_LTC["free"]) + float(balance_LTC["locked"])
        balance_usd_LTC_usd_present = balance_usd_ETH_usd_free_locked_sum * float(symbol_LTCBTC["price"])
        for i in range(1):
                import time
                import yfinance as yf
                real_date = str(time.strftime("%Y-%m-%d"))
                dataETHBTC = yf.download("ETH-BTC", start="2022-09-01", end=real_date, interval="1h")
                dataETHBTC_fc = dataETHBTC.filter(["Close"])
                dataETHBTC_fc.to_csv("ETHBTC_1h.csv")

                print(dataETHBTC_fc)
                plt.plot(dataETHBTC_fc)
                plt.savefig('data_ETHBTC_filter_close_1m.jpg')

                import os
                import math
                import numpy as np
                from sklearn.preprocessing import MinMaxScaler
                from tensorflow.keras.models import load_model
                from tensorflow.keras.models import save_model
                from tensorflow.keras.models import Sequential
                from tensorflow.keras.layers import Dense, LSTM
                import matplotlib.pyplot as plt

                plt.style.use('fivethirtyeight')

                start_time = time.time()
                symbol_ETHBTC = info[0]
                symbol_ETHBTC_symbol = symbol_ETHBTC["symbol"]
                symbol_ETHBTC_price = symbol_ETHBTC["price"]
                data_safe_file_ETHBTC = open("ETHBTC.csv","a")
                data_safe_file_ETHBTC.write(str(time.strftime("%Y-%m-%d %H:%M:%S+00:00")))
                data_safe_file_ETHBTC.write(",")
                data_safe_file_ETHBTC.write(str(symbol_ETHBTC_price))
                data_safe_file_ETHBTC.write(",")
                data_safe_file_ETHBTC.write(str(balance_ETH["free"]))
                data_safe_file_ETHBTC.write(",")
                data_safe_file_ETHBTC.write(str(balance_btc["free"]))
                data_safe_file_ETHBTC.write(",")
                data_safe_file_ETHBTC.write("\n")
                data_safe_file_ETHBTC.close()

                data_read_pandas_ETHBTC = pd.read_csv("ETHBTC.csv")
                #data_read_pandas_ETHBTC = data_read_pandas_ETHBTC.tail(500)
                data_read_pandas_ETHBTC_shape_row,data_read_pandas_ETHBTC_shape_col = data_read_pandas_ETHBTC.shape[0],data_read_pandas_ETHBTC.shape[1]
                print(data_read_pandas_ETHBTC.shape)
                print([data_read_pandas_ETHBTC_shape_row,data_read_pandas_ETHBTC_shape_col])

                filter_ETHBTC_price = data_read_pandas_ETHBTC.filter(["price_ETHBTC"])

                print(filter_ETHBTC_price)

                # create dATEFRAME CLOSE
                data = data_read_pandas_ETHBTC.filter(["price_ETHBTC"])

                # data_df_pandas_filter = data_df_pandas.filter(["Well"])
                print(data)

                # convert dataframe
                dataset = data.values

                # dataset  = data_df_pandas_filter.values
                print(dataset)

                # get the number rows to train the model
                training_data_len = math.ceil(len(dataset) * .8)
                print(training_data_len)
                # scale the data
                scaler = MinMaxScaler(feature_range=(0, 1))
                scaled_data = scaler.fit_transform(dataset)
                print(scaled_data)
                plt.plot(scaled_data)
                plt.savefig("scaled_data_ETHBTC.png")

                # create the training dataset
                train_data = scaled_data[0:training_data_len, :]
                # split the data into x_train and y_train data sets
                x_train = []
                y_train = []
                for rar in range(60, len(train_data)):
                    x_train.append(train_data[rar - 60:rar, 0])
                    y_train.append(train_data[rar, 0])
                    if rar <= 61:
                        print(x_train)
                        print(y_train)
                        print()
                # conver the x_train and y_train to numpy arrays
                x_train, y_train = np.array(x_train), np.array(y_train)
                # reshape the data
                x_train = np.reshape(x_train, (x_train.shape[0], x_train.shape[1], 1))
                print(x_train.shape)
                import tensorflow as tf

                # biuld to LST model

                model = Sequential()
                model.add(LSTM(50, return_sequences=True, input_shape=(x_train.shape[1], 1)))
                model.add(LSTM(101, return_sequences=False))
                model.add(Dense(50))
                model.add(Dense(25))
                model.add(Dense(1))
                # cmopale th emodel
                model.compile(optimizer='adam', loss='mean_squared_error')
                # train_the_model
                model.summary()
                print("Fit model on training data")

                # Evaluate the model on the test data using `evaluate`
                print("Evaluate on test data")
                results = model.evaluate(x_train, y_train, batch_size=1)
                print("test loss, test acc:", results)

                model = tf.keras.models.load_model(os.path.join("./dnn/", "ETHBTC_model.h5"))
                model.fit(x_train, y_train, batch_size=1, epochs=1)

                model.save(os.path.join("./dnn/", "ETHBTC_model.h5"))
                #reconstructed_model = tf.keras.models.load_model(os.path.join("./dnn/", "BTC-RUB_model.h5"))

                #np.testing.assert_allclose(model.predict(x_train), reconstructed_model.predict(x_train))
                #reconstructed_model.fit(x_train, y_train)

                # create the testing data set
                # create a new array containing scaled values from index 1713 to 2216
                test_data = scaled_data[training_data_len - 60:, :]
                # create the fata sets x_test and y_test
                x_test = []
                y_test = dataset[training_data_len:, :]
                for resr in range(60, len(test_data)):
                    x_test.append(test_data[resr - 60:resr, 0])

                # conert the data to numpy array
                x_test = np.array(x_test)

                # reshape the data
                x_test = np.reshape(x_test, (x_test.shape[0], x_test.shape[1], 1))

                # get the model predicted price values
                predictions = model.predict(x_test)
                predictions = scaler.inverse_transform(predictions)

                # get the root squared error (RMSE)
                rmse = np.sqrt(np.mean(predictions - y_test) ** 2)
                print(rmse)
                if int(rmse) >= 1000:
                    from termcolor import colored

                    print(colored("Высокий уровень ошибки", "red"))

                # get the quate

                new_df = data_read_pandas_ETHBTC.filter(["price_ETHBTC"])

                # get teh last 60 days closing price values and convert the dataframe to an array
                last_60_days = new_df[-60:].values
                # scale the data to be values beatwet 0 and 1

                last_60_days_scaled = scaler.transform(last_60_days)

                # creAte an enemy list
                X_test = []
                # Append past 60 days
                X_test.append(last_60_days_scaled)

                # convert the x tesst dataset to numpy
                X_test = np.array(X_test)

                # Reshape the dataframe
                X_test = np.reshape(X_test, (X_test.shape[0], X_test.shape[1], 1))
                # get predict scaled

                pred_price = model.predict(X_test)
                # undo the scaling
                pred_price = scaler.inverse_transform(pred_price)
                print(pred_price)

                pred_price_a = pred_price[0]
                pred_price_aa = pred_price_a[0]
                preset_pred_price = round(pred_price_aa,6)
                print(pred_price)
                print(preset_pred_price)
                old_time = time.time() - start_time
                print("Время на расчеты :" + str(old_time))
                min_lovume_btc = 0.00011

                #info = client.get_symbol_info('ETHBTC')
                #print(info)
                #print(info['filters'][2]['minQty'])


                time.sleep(5)

                #pred_price = float(symbol_BTCRUB_price) + float(random.randint(-1000,1000))

                if preset_pred_price <= float(symbol_ETHBTC_price):
                    info = client.get_all_tickers()
                    symbol_ETHBTC = info[0]
                    a = float(1)
                    b = float(balance_btc["free"])
                    ab_sum = a * b
                    data_coin = float(ab_sum) - min_lovume_btc
                    print(data_coin)

                    if data_coin <= 0:
                        print([data_coin, a, b])
                        print(ab_sum)
                        quantity = float(min_lovume_btc / float(symbol_ETHBTC_price))
                        print(quantity)
                        print("Недостаточно  btc")
                    elif data_coin >= 0:
                        print([data_coin, a, b])
                        print("\n" + "SELL Покупать btc  " + str(preset_pred_price))
                        print(a)
                        quantity = float(min_lovume_btc / float(symbol_ETHBTC_price))
                        quantity_start = round(quantity, 4)
                        print(quantity_start)
                        order = client.order_limit_buy(symbol='ETHBTC',quantity=quantity_start,price=preset_pred_price)
                        print(order)
                        data_safe_file_BTCRUB = open("ETHBTCorder.csv", "a")
                        data_safe_file_BTCRUB.write(str(time.strftime("%Y-%m-%d %H:%M:%S+00:00")))
                        data_safe_file_BTCRUB.write(",")
                        data_safe_file_BTCRUB.write(str(symbol_ETHBTC_price))
                        data_safe_file_BTCRUB.write(",")
                        data_safe_file_BTCRUB.write(str(order))
                        data_safe_file_BTCRUB.write(",")
                        data_safe_file_BTCRUB.write("\n")
                        data_safe_file_BTCRUB.close()



                elif preset_pred_price >= float(symbol_ETHBTC_price):
                    info = client.get_all_tickers()
                    symbol_ETHBTC = info[0]
                    a = float(symbol_ETHBTC_price)
                    b = float(balance_ETH["free"])
                    ab_sum = a * b
                    data_coin = float(ab_sum) - min_lovume_btc
                    print(data_coin)

                    if data_coin <= 0:
                        print([data_coin, a, b])
                        print(ab_sum)
                        quantity = float(min_lovume_btc / float(symbol_ETHBTC_price))
                        print(quantity)
                        print("Недостаточно BTC для продажи")
                    elif data_coin >= 0:
                        print([data_coin, a, b])
                        print("\n" + "BUY Покупать за BTC " + str(preset_pred_price))
                        print(a)
                        quantity = float(min_lovume_btc / float(symbol_ETHBTC_price))
                        quantity_start = round(quantity, 4)
                        print(quantity_start)
                        order = client.order_limit_sell(symbol='ETHBTC',quantity=quantity_start,price=preset_pred_price)
                        print(order)
                        data_safe_file_ETHBTC = open("ETHBTCorder.csv", "a")
                        data_safe_file_ETHBTC.write(str(time.strftime("%Y-%m-%d %H:%M:%S+00:00")))
                        data_safe_file_ETHBTC.write(",")
                        data_safe_file_ETHBTC.write(str(symbol_ETHBTC_price))
                        data_safe_file_ETHBTC.write(",")
                        data_safe_file_ETHBTC.write(str(order))
                        data_safe_file_ETHBTC.write(",")
                        data_safe_file_ETHBTC.write("\n")
                        data_safe_file_ETHBTC.close()
                
